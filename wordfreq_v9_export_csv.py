#!/usr/bin/env python3
# -*- coding: utf-8 -*-

#from operator import itemgetter
from docx.api import Document
from urllib import parse,request 
import re
import csv

#change the different tables for program
#table = document.tables[0]
#    z = [0]
p = re.compile('[0-9.]+')
n = re.compile('<TD>  .*\(Na\) ')
t = re.compile(': .*\(Na\)')
c = re.compile('Word not found')

#Import word document
document = Document('yyy_items_B.docx')
#for o in range(48):
table = document.tables[0]
data = []

keys = None
for i, row in enumerate(table.rows):
    text = (cell.text.replace('\n','') for cell in row.cells)

    if i == 0:
        keys = tuple(text)
        continue
    row_data = dict(zip(keys, text))
    data.append(row_data)


#print(data)
#data = sorted(data, key=itemgetter('排序'))

#get item name to input to website
items = []
for entry in data:
    a = entry['反應項']
    items.append(a)
#print(items)
items_na = [item + "(Nc)" for item in items]


#Website
URL = 'http://elearning.ling.sinica.edu.tw/Cfindwordfreq.php'

dat = []
ke = ('類別','yyy排序','代號','AS排名','頻率','百分比','累積百分比','反應項')

for j in range(len(items_na)):
    i = items_na[j]
    word = i.encode('big5')
# word need encoding first, check website encoding
    q = parse.urlencode({'wordforfreq':word}).encode('big5')
    h = {'User-Agent':'Mozilla/5.0'} 
    req = request.Request(URL,q,h)
    datacontent = request.urlopen(req).read()
    webdata = datacontent.decode('big5','ignore')
    numbers = p.findall(webdata)
    noexist = c.findall(webdata)
    naexist = n.findall(webdata)
    nexist = t.findall(webdata)
    z = [0]
    j = [j+1]

#different situations: 
    if(noexist):
        j.extend(['na']*5)
    
        if(nexist):
            nexist = [nexist[0].replace(': ','')]
            j.extend(nexist)
            z.extend(j)
            zipdict = dict(zip(ke, z))
            dat.append(zipdict)

    elif(numbers):
        j.extend(numbers[1:6])
        
        if(naexist):
            naexist = [naexist[0].replace(' ','').replace('<TD>','')]
            j.extend(naexist)
            z.extend(j)
            zipdict = dict(zip(ke, z))
            dat.append(zipdict)
#Done:print(dat) to check

#Combine website and word doc contents
for i in range(len(data)):
    data[i].update(dat[i])
    
#output as csv
with open('itemrank_39_city.csv', 'a', newline='',encoding='big5') as csvfile:
    fieldnames = list(data[0].keys())
    writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
#    if o <1:
    writer.writeheader()
#        else:
#            break
    for row in data:
        writer.writerow(row)



#dat = sorted(dat, key=itemgetter('AS排名'))

